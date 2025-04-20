
from flask import Flask, request, send_file
from docx import Document
import os
import io
import tempfile
from flask_cors import CORS

app = Flask(__name__)
CORS(app)

def format_date(d):
    return d.strip() if d else ""

def replace_placeholders(doc, placeholders):
    for para in doc.paragraphs:
        for run in para.runs:
            for key, val in placeholders.items():
                if key in run.text:
                    run.text = run.text.replace(key, val)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        for key, val in placeholders.items():
                            if key in run.text:
                                run.text = run.text.replace(key, val)

@app.route('/')
def index():
    return "DOCX contract generator for GAS - DOMACNOST is running."

@app.route('/api/generate', methods=['POST'])
def generate():
    data = request.json

    doc = Document("Rekapitulace_Domacnost_Plyn.docx")
    placeholders = {
        "{{cislo_smlouvy}}": data.get("cislo_smlouvy", ""),
        "{{cislo_partnera}}": data.get("cislo_partnera", ""),
        "{{jmeno}}": data.get("jmeno", ""),
        "{{prijmeni}}": data.get("prijmeni", ""),
        "{{datum_narozeni}}": format_date(data.get("datum_narozeni", "")),
        "{{ulice_trvala}}": data.get("ulice_trvala", ""),
        "{{mesto_trvala}}": data.get("mesto_trvala", ""),
        "{{psc_trvala}}": data.get("psc_trvala", ""),
        "{{email}}": data.get("email", ""),
        "{{telefon}}": data.get("telefon", ""),
        "{{zpusob_odesilani}}": data.get("zpusob_odesilani", ""),
        "{{platby_faktury}}": data.get("platby_faktury", ""),
        "{{platby_zalohy}}": data.get("platby_zalohy", ""),
        "{{cislo_uctu}}": data.get("cislo_uctu", ""),
        "{{zahajeni_dodavek}}": format_date(data.get("zahajeni_dodavek", "")),
        "{{prolongace}}": format_date(data.get("prolongace", "")),
        "{{ean}}": data.get("ean", ""),
        "{{ulice_odber}}": data.get("ulice_odber", ""),
        "{{mesto_odber}}": data.get("mesto_odber", ""),
        "{{psc_odber}}": data.get("psc_odber", "")
    }

    replace_placeholders(doc, placeholders)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        tmp.seek(0)
        return send_file(tmp.name, as_attachment=True, download_name="Rekapitulace_smlouvy_plyn_domacnost.docx")

if __name__ == "__main__":
    app.run(debug=False, host="0.0.0.0", port=10000)
